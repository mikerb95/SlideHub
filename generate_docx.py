import csv
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def main():
    doc = Document()
    
    # Title
    heading = doc.add_heading('Historias de Usuario - SlideHub', 0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Read CSV
    with open('docs/Historias de Usuario - SlideHub.csv', mode='r', encoding='utf-8-sig') as f:
        # Use comma as separator
        reader = csv.reader(f)
        rows = list(reader)

    # Some CSVs may have a different delimiter or encoding, but based on the printout, they are comma-separated and double-quoted.
    # Start reading from row 3 (skip headers)
    
    current_story = None
    
    for i, row in enumerate(rows):
        if i < 3:
            continue
            
        if not row or len(row) < 10:
            continue
        
        # Determine if it's a new story or just another scenario
        story_id = row[1].strip()
        
        if story_id:
            # It's a new story
            doc.add_heading(f"[{story_id}] {row[3].strip()}", level=1)
            
            p = doc.add_paragraph()
            p.add_run("Rol:").bold = True
            p.add_run(f" {row[2].strip()}\n")
            
            p.add_run("Razón / Resultado:").bold = True
            p.add_run(f" {row[4].strip()}\n")
            
            doc.add_heading("Criterios de Aceptación", level=2)
            
        scenario_num = row[5].strip()
        scenario_title = row[6].strip()
        context = row[7].strip()
        event = row[8].strip()
        expected = row[9].strip()
        
        if scenario_num:
            doc.add_heading(f"Escenario {scenario_num}: {scenario_title}", level=3)
            
            p_crit = doc.add_paragraph()
            p_crit.add_run("Dado (Contexto):").bold = True
            p_crit.add_run(f" {context}\n")
            
            p_crit.add_run("Cuando (Evento):").bold = True
            p_crit.add_run(f" {event}\n")
            
            p_crit.add_run("Entonces (Resultado):").bold = True
            p_crit.add_run(f" {expected}\n")

    output_path = 'docs/Historias_de_Usuario_SlideHub.docx'
    doc.save(output_path)
    print(f"Generated {output_path}")

if __name__ == '__main__':
    main()
